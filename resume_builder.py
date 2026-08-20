import asyncio
import base64
import html
import io

import streamlit as st

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import (
    Image,
    Paragraph,
    SimpleDocTemplate,
    Table,
    TableStyle,
)


# ============================================================
# OPTIONAL PLAYWRIGHT
# ============================================================

try:
    from playwright.async_api import async_playwright

    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False


# ============================================================
# CONSTANTS
# ============================================================

A4_WIDTH_PX = 794
A4_HEIGHT_PX = 1123

DEFAULT_AVATAR = (
    "https://images.unsplash.com/"
    "photo-1573496359142-b8d87734a5a2"
    "?auto=format&fit=crop&q=80&w=300"
)


# ============================================================
# SMALL HELPERS
# ============================================================

def safe(value):
    """Safely convert user input to HTML."""
    if value is None:
        return ""

    return html.escape(str(value), quote=True)


def clean_list(value):
    """Normalize list-like values."""
    if value is None:
        return []

    if isinstance(value, list):
        return [
            str(item).strip()
            for item in value
            if str(item).strip()
        ]

    return [
        item.strip()
        for item in str(value).split(",")
        if item.strip()
    ]


def set_cell_background(cell, fill_hex):
    tc_pr = cell._element.get_or_add_tcPr()

    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    )

    tc_pr.append(shd)


def remove_table_borders(table):
    tbl_pr = table._element.xpath("w:tblPr")

    if tbl_pr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:bottom w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="none"/>'
            f'<w:insideV w:val="none"/>'
            f"</w:tblBorders>"
        )

        tbl_pr[0].append(borders)


# ============================================================
# PLAYWRIGHT HTML → PDF
# ============================================================

async def _render_html_to_pdf_async(html_content):
    async with async_playwright() as p:

        browser = await p.chromium.launch(headless=True)

        page = await browser.new_page(
            viewport={
                "width": A4_WIDTH_PX,
                "height": A4_HEIGHT_PX,
            },
            device_scale_factor=1,
        )

        await page.set_content(
            html_content,
            wait_until="networkidle",
        )

        # Give fonts / layout a moment to settle.
        await page.wait_for_timeout(500)

        pdf_bytes = await page.pdf(
            format="A4",
            print_background=True,
            prefer_css_page_size=True,
            margin={
                "top": "0",
                "right": "0",
                "bottom": "0",
                "left": "0",
            },
        )

        await browser.close()

        return pdf_bytes


def generate_pdf_from_html(html_content):
    if not PLAYWRIGHT_AVAILABLE:
        return None

    try:
        return asyncio.run(
            _render_html_to_pdf_async(html_content)
        )

    except Exception:
        return None


# ============================================================
# REPORTLAB FALLBACK PDF
# ============================================================

def generate_pdf_bytes(user_profile, image_bytes=None):

    buffer = io.BytesIO()

    page_width, page_height = A4

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36,
    )

    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "HeaderName",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=26,
        leading=31,
        textColor=colors.white,
        spaceAfter=5,
    )

    title_style = ParagraphStyle(
        "HeaderTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        textColor=colors.HexColor("#38BDF8"),
    )

    side_heading = ParagraphStyle(
        "SideHeading",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor("#0F172A"),
        spaceBefore=12,
        spaceAfter=6,
    )

    side_body = ParagraphStyle(
        "SideBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=14,
        textColor=colors.HexColor("#334155"),
    )

    main_heading = ParagraphStyle(
        "MainHeading",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#0F172A"),
        spaceBefore=14,
        spaceAfter=6,
    )

    main_body = ParagraphStyle(
        "MainBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.7,
        leading=15,
        textColor=colors.HexColor("#334155"),
    )

    main_sub = ParagraphStyle(
        "MainSub",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9.7,
        leading=13,
        textColor=colors.HexColor("#0284C7"),
    )

    dates_style = ParagraphStyle(
        "DatesStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#64748B"),
        spaceAfter=5,
    )

    name = (
        user_profile.get("name", "").strip()
        or "Galena Micheal"
    )

    role = (
        user_profile.get("current_role", "").strip().upper()
        or "BACHELOR OF ARTS IN EDUCATION"
    )

    # --------------------------------------------------------
    # HEADER
    # --------------------------------------------------------

    avatar = []

    if image_bytes:
        try:
            avatar.append(
                Image(
                    io.BytesIO(image_bytes),
                    width=72,
                    height=72,
                )
            )
        except Exception:
            pass

    header_text = [
        Paragraph(
            safe(name),
            name_style,
        ),
        Paragraph(
            safe(role),
            title_style,
        ),
    ]

    header_table = Table(
        [[avatar, header_text]],
        colWidths=[
            page_width * 0.30,
            page_width * 0.70,
        ],
    )

    header_table.setStyle(
        TableStyle(
            [
                (
                    "BACKGROUND",
                    (0, 0),
                    (-1, -1),
                    colors.HexColor("#0F172A"),
                ),
                (
                    "VALIGN",
                    (0, 0),
                    (-1, -1),
                    "MIDDLE",
                ),
                (
                    "TOPPADDING",
                    (0, 0),
                    (-1, -1),
                    24,
                ),
                (
                    "BOTTOMPADDING",
                    (0, 0),
                    (-1, -1),
                    24,
                ),
                (
                    "LEFTPADDING",
                    (0, 0),
                    (0, 0),
                    22,
                ),
                (
                    "LEFTPADDING",
                    (1, 0),
                    (1, 0),
                    14,
                ),
                (
                    "RIGHTPADDING",
                    (1, 0),
                    (1, 0),
                    24,
                ),
            ]
        )
    )

    # --------------------------------------------------------
    # SIDEBAR
    # --------------------------------------------------------

    sidebar = []

    contact_items = [
        str(user_profile.get(key, "")).strip()
        for key in [
            "address",
            "phone",
            "email",
            "website",
        ]
        if str(user_profile.get(key, "")).strip()
    ]

    if contact_items:
        sidebar.append(
            Paragraph(
                "CONTACT",
                side_heading,
            )
        )

        for item in contact_items:
            sidebar.append(
                Paragraph(
                    safe(item),
                    side_body,
                )
            )

    languages = clean_list(
        user_profile.get("languages", [])
    )

    if languages:
        sidebar.append(
            Paragraph(
                "LANGUAGES",
                side_heading,
            )
        )

        for language in languages:
            sidebar.append(
                Paragraph(
                    f"• {safe(language)}",
                    side_body,
                )
            )

    skills = clean_list(
        user_profile.get("skills", [])
    )

    if skills:
        sidebar.append(
            Paragraph(
                "SKILLS",
                side_heading,
            )
        )

        for skill in skills:
            sidebar.append(
                Paragraph(
                    f"✓ {safe(skill)}",
                    side_body,
                )
            )

    # --------------------------------------------------------
    # MAIN CONTENT
    # --------------------------------------------------------

    main = []

    summary = (
        user_profile.get("summary", "")
        .strip()
    )

    if summary:
        main.append(
            Paragraph(
                "SUMMARY",
                main_heading,
            )
        )

        main.append(
            Paragraph(
                safe(summary),
                main_body,
            )
        )

    institution = (
        user_profile.get(
            "edu_institution",
            "",
        ).strip()
    )

    if institution:

        main.append(
            Paragraph(
                "EDUCATION",
                main_heading,
            )
        )

        location = (
            user_profile.get(
                "edu_location",
                "",
            ).strip()
        )

        degree = (
            user_profile.get(
                "edu_degree",
                "",
            ).strip()
        )

        dates = (
            user_profile.get(
                "edu_dates",
                "",
            ).strip()
        )

        institution_text = (
            f"<b>{safe(institution)}</b>"
        )

        if location:
            institution_text += (
                f", {safe(location)}"
            )

        main.append(
            Paragraph(
                institution_text,
                main_body,
            )
        )

        if degree:
            main.append(
                Paragraph(
                    safe(degree),
                    main_sub,
                )
            )

        if dates:
            main.append(
                Paragraph(
                    safe(dates),
                    dates_style,
                )
            )

    company = (
        user_profile.get(
            "exp_company",
            "",
        ).strip()
    )

    if company:

        main.append(
            Paragraph(
                "EXPERIENCE",
                main_heading,
            )
        )

        role_value = (
            user_profile.get(
                "exp_role",
                "",
            ).strip()
        )

        dates = (
            user_profile.get(
                "exp_dates",
                "",
            ).strip()
        )

        company_text = (
            f"<b>{safe(company)}</b>"
        )

        if role_value:
            company_text += (
                f", <i>{safe(role_value)}</i>"
            )

        main.append(
            Paragraph(
                company_text,
                main_body,
            )
        )

        if dates:
            main.append(
                Paragraph(
                    safe(dates),
                    dates_style,
                )
            )

        bullets = clean_list(
            user_profile.get(
                "exp_bullets",
                [],
            )
        )

        for bullet in bullets:
            main.append(
                Paragraph(
                    f"▪ {safe(bullet)}",
                    main_body,
                )
            )

    certifications = clean_list(
        user_profile.get(
            "certifications",
            [],
        )
    )

    if certifications:

        main.append(
            Paragraph(
                "CERTIFICATIONS",
                main_heading,
            )
        )

        for certification in certifications:
            main.append(
                Paragraph(
                    f"▪ {safe(certification)}",
                    main_body,
                )
            )

    body_table = Table(
        [[sidebar, main]],
        colWidths=[
            page_width * 0.30,
            page_width * 0.70,
        ],
    )

    body_table.setStyle(
        TableStyle(
            [
                (
                    "BACKGROUND",
                    (0, 0),
                    (0, 0),
                    colors.HexColor("#F8FAFC"),
                ),
                (
                    "BACKGROUND",
                    (1, 0),
                    (1, 0),
                    colors.white,
                ),
                (
                    "VALIGN",
                    (0, 0),
                    (-1, -1),
                    "TOP",
                ),
                (
                    "TOPPADDING",
                    (0, 0),
                    (-1, -1),
                    22,
                ),
                (
                    "BOTTOMPADDING",
                    (0, 0),
                    (-1, -1),
                    22,
                ),
                (
                    "LEFTPADDING",
                    (0, 0),
                    (0, 0),
                    20,
                ),
                (
                    "RIGHTPADDING",
                    (0, 0),
                    (0, 0),
                    16,
                ),
                (
                    "LEFTPADDING",
                    (1, 0),
                    (1, 0),
                    26,
                ),
                (
                    "RIGHTPADDING",
                    (1, 0),
                    (1, 0),
                    26,
                ),
            ]
        )
    )

    doc.build(
        [
            header_table,
            body_table,
        ]
    )

    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# DOCX GENERATOR
# ============================================================

def generate_docx_bytes(
    user_profile,
    image_bytes=None,
):

    doc = Document()

    section = doc.sections[0]

    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

    # --------------------------------------------------------
    # HEADER
    # --------------------------------------------------------

    header = doc.add_table(
        rows=1,
        cols=2,
    )

    header.autofit = False
    header.alignment = WD_TABLE_ALIGNMENT.CENTER

    remove_table_borders(header)

    image_cell = header.rows[0].cells[0]
    text_cell = header.rows[0].cells[1]

    image_cell.width = Inches(2.5)
    text_cell.width = Inches(5.77)

    set_cell_background(
        image_cell,
        "0F172A",
    )

    set_cell_background(
        text_cell,
        "0F172A",
    )

    image_cell.vertical_alignment = (
        WD_ALIGN_VERTICAL.CENTER
    )

    text_cell.vertical_alignment = (
        WD_ALIGN_VERTICAL.CENTER
    )

    if image_bytes:
        try:
            p = image_cell.paragraphs[0]

            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            run = p.add_run()

            run.add_picture(
                io.BytesIO(image_bytes),
                width=Inches(1.05),
            )

        except Exception:
            pass

    name = (
        user_profile.get("name", "").strip()
        or "Galena Micheal"
    )

    role = (
        user_profile.get(
            "current_role",
            "",
        ).strip().upper()
        or "BACHELOR OF ARTS IN EDUCATION"
    )

    p_name = text_cell.paragraphs[0]

    p_name.paragraph_format.space_before = Pt(15)
    p_name.paragraph_format.space_after = Pt(3)

    run_name = p_name.add_run(name)

    run_name.bold = True
    run_name.font.size = Pt(25)
    run_name.font.color.rgb = RGBColor(
        255,
        255,
        255,
    )

    p_title = text_cell.add_paragraph()

    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(15)

    run_title = p_title.add_run(role)

    run_title.bold = True
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGBColor(
        56,
        189,
        248,
    )

    # --------------------------------------------------------
    # BODY
    # --------------------------------------------------------

    body = doc.add_table(
        rows=1,
        cols=2,
    )

    body.autofit = False
    body.alignment = WD_TABLE_ALIGNMENT.CENTER

    remove_table_borders(body)

    left = body.rows[0].cells[0]
    right = body.rows[0].cells[1]

    left.width = Inches(2.5)
    right.width = Inches(5.77)

    set_cell_background(
        left,
        "F8FAFC",
    )

    set_cell_background(
        right,
        "FFFFFF",
    )

    left.vertical_alignment = (
        WD_ALIGN_VERTICAL.TOP
    )

    right.vertical_alignment = (
        WD_ALIGN_VERTICAL.TOP
    )

    def sidebar_heading(
        cell,
        text,
    ):
        p = cell.add_paragraph()

        p.paragraph_format.space_before = Pt(15)
        p.paragraph_format.space_after = Pt(5)

        r = p.add_run(text)

        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(
            15,
            23,
            42,
        )

    def main_heading(
        cell,
        text,
    ):
        p = cell.add_paragraph()

        p.paragraph_format.space_before = Pt(17)
        p.paragraph_format.space_after = Pt(6)

        r = p.add_run(text)

        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(
            15,
            23,
            42,
        )

    def body_text(
        cell,
        text,
        size=9.7,
    ):
        p = cell.add_paragraph()

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)

        r = p.add_run(text)

        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(
            51,
            65,
            85,
        )

    # Sidebar

    contact = [
        str(
            user_profile.get(
                key,
                "",
            )
        ).strip()
        for key in [
            "address",
            "phone",
            "email",
            "website",
        ]
        if str(
            user_profile.get(
                key,
                "",
            )
        ).strip()
    ]

    if contact:

        sidebar_heading(
            left,
            "CONTACT",
        )

        for item in contact:
            body_text(
                left,
                item,
            )

    languages = clean_list(
        user_profile.get(
            "languages",
            [],
        )
    )

    if languages:

        sidebar_heading(
            left,
            "LANGUAGES",
        )

        for item in languages:
            body_text(
                left,
                f"• {item}",
            )

    skills = clean_list(
        user_profile.get(
            "skills",
            [],
        )
    )

    if skills:

        sidebar_heading(
            left,
            "SKILLS",
        )

        for item in skills:
            body_text(
                left,
                f"✓ {item}",
            )

    # Main

    summary = (
        user_profile.get(
            "summary",
            "",
        ).strip()
    )

    if summary:

        main_heading(
            right,
            "SUMMARY",
        )

        body_text(
            right,
            summary,
        )

    institution = (
        user_profile.get(
            "edu_institution",
            "",
        ).strip()
    )

    if institution:

        main_heading(
            right,
            "EDUCATION",
        )

        location = (
            user_profile.get(
                "edu_location",
                "",
            ).strip()
        )

        degree = (
            user_profile.get(
                "edu_degree",
                "",
            ).strip()
        )

        dates = (
            user_profile.get(
                "edu_dates",
                "",
            ).strip()
        )

        body_text(
            right,
            f"{institution}"
            + (
                f", {location}"
                if location
                else ""
            ),
            size=9,
        )

        if degree:
            p = right.add_paragraph()

            p.paragraph_format.space_after = Pt(3)

            r = p.add_run(degree)

            r.italic = True
            r.font.size = Pt(8.7)
            r.font.color.rgb = RGBColor(
                2,
                132,
                199,
            )

        if dates:
            body_text(
                right,
                dates,
                size=7.5,
            )

    company = (
        user_profile.get(
            "exp_company",
            "",
        ).strip()
    )

    if company:

        main_heading(
            right,
            "EXPERIENCE",
        )

        role_value = (
            user_profile.get(
                "exp_role",
                "",
            ).strip()
        )

        company_text = company

        if role_value:
            company_text += (
                f", {role_value}"
            )

        body_text(
            right,
            company_text,
            size=10,
        )

        dates = (
            user_profile.get(
                "exp_dates",
                "",
            ).strip()
        )

        if dates:
            body_text(
                right,
                dates,
                size=8.5,
            )

        bullets = clean_list(
            user_profile.get(
                "exp_bullets",
                [],
            )
        )

        for bullet in bullets:
            body_text(
                right,
                f"▪ {bullet}",
            )

    certifications = clean_list(
        user_profile.get(
            "certifications",
            [],
        )
    )

    if certifications:

        main_heading(
            right,
            "CERTIFICATIONS",
        )

        for certification in certifications:
            body_text(
                right,
                f"▪ {certification}",
            )

    buffer = io.BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# TAILWIND RESUME HTML
# ============================================================

def build_resume_html(
    name,
    professional_title,
    address,
    phone,
    email,
    website,
    languages,
    skills,
    summary,
    edu_institution,
    edu_location,
    edu_degree,
    edu_dates,
    exp_company,
    exp_role,
    exp_dates,
    experience_bullets,
    certifications,
    avatar_src,
):

    name_display = (
        name.strip()
        or "Galena Micheal"
    )

    title_display = (
        professional_title.strip().upper()
        or "BACHELOR OF ARTS IN EDUCATION"
    )

    contact_items = [
        address,
        phone,
        email,
        website,
    ]

    contact_items = [
        safe(item.strip())
        for item in contact_items
        if item.strip()
    ]

    contact_html = "".join(
        f"""
        <div class="mb-2.5 break-words">
            {item}
        </div>
        """
        for item in contact_items
    )

    if not contact_html:
        contact_html = """
        <span class="text-slate-400 italic">
            Contact information
        </span>
        """

    language_html = "".join(
        f"""
        <div class="mb-2 text-[13px] text-slate-600">
            <span class="mr-2 text-sky-600">•</span>
            {safe(language)}
        </div>
        """
        for language in languages
    )

    if not language_html:
        language_html = """
        <span class="text-slate-400 italic text-xs">
            No languages
        </span>
        """

    skill_html = "".join(
        f"""
        <div class="mb-2.5 flex items-start text-[12px] leading-[1.45] text-slate-600">
            <span class="mr-2 mt-[1px] font-bold text-sky-600">
                ✓
            </span>

            <span>
                {safe(skill)}
            </span>
        </div>
        """
        for skill in skills
    )

    if not skill_html:
        skill_html = """
        <span class="text-slate-400 italic text-xs">
            No skills
        </span>
        """

    if summary.strip():
        summary_html = f"""
        <p class="m-0 text-[13px] leading-[1.65] text-slate-600">
            {safe(summary.strip())}
        </p>
        """
    else:
        summary_html = """
        <p class="m-0 text-xs italic text-slate-400">
            Professional summary
        </p>
        """

    # --------------------------------------------------------
    # EDUCATION
    # --------------------------------------------------------

    if any(
        [
            edu_institution.strip(),
            edu_location.strip(),
            edu_degree.strip(),
            edu_dates.strip(),
        ]
    ):

        location_html = ""

        if edu_location.strip():
            location_html = (
                f", {safe(edu_location)}"
            )

        edu_html = f"""
        <div class="text-[14px] font-bold text-slate-900">
            {safe(edu_institution)}
            {location_html}
        </div>
        """

        if edu_degree.strip():
            edu_html += f"""
            <div class="mt-1 text-[12.5px] italic text-sky-600">
                {safe(edu_degree)}
            </div>
            """

        if edu_dates.strip():
            edu_html += f"""
            <div class="mt-1 text-[12px] font-semibold text-slate-500">
                {safe(edu_dates)}
            </div>
            """

    else:

        edu_html = """
        <p class="m-0 text-xs italic text-slate-400">
            Education
        </p>
        """

    # --------------------------------------------------------
    # EXPERIENCE
    # --------------------------------------------------------

    if any(
        [
            exp_company.strip(),
            exp_role.strip(),
            exp_dates.strip(),
            experience_bullets,
        ]
    ):

        role_html = ""

        if exp_role.strip():
            role_html = f"""
            <span class="italic font-normal text-slate-600">
                , {safe(exp_role)}
            </span>
            """

        exp_html = f"""
        <div class="text-[14px] text-slate-900">
            <span class="font-bold">
                {safe(exp_company)}
            </span>
            {role_html}
        </div>
        """

        if exp_dates.strip():
            exp_html += f"""
            <div class="mt-1 mb-2.5 text-[12px] font-semibold text-slate-500">
                {safe(exp_dates)}
            </div>
            """

        if experience_bullets:

            bullets_html = "".join(
                f"""
                <li class="mb-2 pl-1">
                    {safe(bullet)}
                </li>
                """
                for bullet in experience_bullets
            )

            exp_html += f"""
            <ul class="m-0 list-square pl-5 text-[13px] leading-[1.55] text-slate-600">
                {bullets_html}
            </ul>
            """

    else:

        exp_html = """
        <p class="m-0 text-xs italic text-slate-400">
            Work experience
        </p>
        """

    # --------------------------------------------------------
    # CERTIFICATIONS
    # --------------------------------------------------------

    if certifications:

        certification_html = "".join(
            f"""
            <li class="mb-2 pl-1">
                {safe(certification)}
            </li>
            """
            for certification in certifications
        )

        cert_html = f"""
        <ul class="m-0 list-square pl-5 text-[13px] leading-[1.55] text-slate-600">
            {certification_html}
        </ul>
        """

    else:

        cert_html = """
        <p class="m-0 text-xs italic text-slate-400">
            Certifications
        </p>
        """

    # --------------------------------------------------------
    # FULL HTML
    # --------------------------------------------------------

    return f"""
<!DOCTYPE html>

<html lang="en">

<head>

<meta charset="UTF-8">

<meta name="viewport"
      content="width=device-width, initial-scale=1.0">

<script src="https://cdn.tailwindcss.com"></script>

<script>
tailwind.config = {{
    theme: {{
        extend: {{
            fontFamily: {{
                sans: [
                    "Inter",
                    "ui-sans-serif",
                    "system-ui",
                    "sans-serif"
                ],
            }},
        }}
    }}
}}
</script>

<style>

@page {{
    size: A4;
    margin: 0;
}}

html,
body {{
    margin: 0;
    padding: 0;
    background: #e2e8f0;
    font-family: Inter, ui-sans-serif, system-ui, sans-serif;
    -webkit-print-color-adjust: exact !important;
    print-color-adjust: exact !important;
}}

* {{
    box-sizing: border-box;
}}

.resume-page {{
    width: 794px;
    min-height: 1123px;

    margin: 28px auto;

    background: white;

    overflow: hidden;

    box-shadow:
        0 20px 50px rgba(15, 23, 42, 0.12);
}}

.resume-header {{
    min-height: 150px;

    position: relative;

    display: flex;
    align-items: center;

    background:
        linear-gradient(
            135deg,
            #0f172a 0%,
            #172554 100%
        );

    padding:
        34px
        42px
        34px
        178px;
}}

.avatar {{
    position: absolute;

    left: 40px;
    top: 24px;

    width: 112px;
    height: 112px;

    border-radius: 999px;

    object-fit: cover;

    border: 4px solid #38bdf8;

    background: #334155;

    box-shadow:
        0 8px 20px rgba(0, 0, 0, 0.25);
}}

.resume-body {{
    display: grid;

    grid-template-columns:
        30%
        70%;

    align-items: start;
}}

.sidebar {{
    min-height: 973px;

    background: #f8fafc;

    border-right:
        1px solid #e2e8f0;

    padding:
        32px
        24px
        36px
        30px;
}}

.main {{
    padding:
        32px
        38px
        42px
        34px;

    background: white;
}}

.sidebar-section {{
    margin-bottom: 28px;
}}

.main-section {{
    margin-bottom: 27px;
}}

.section-title {{
    display: inline-block;

    margin-bottom: 11px;

    padding-bottom: 4px;

    border-bottom:
        2px solid #38bdf8;

    font-size: 12px;
    line-height: 1;

    font-weight: 800;

    letter-spacing: 0.12em;

    text-transform: uppercase;

    color: #0f172a;
}}

.main-title {{
    display: inline-block;

    margin-bottom: 11px;

    padding-bottom: 4px;

    border-bottom:
        2px solid #0f172a;

    font-size: 12px;
    line-height: 1;

    font-weight: 800;

    letter-spacing: 0.12em;

    text-transform: uppercase;

    color: #0f172a;
}}

@media print {{

    html,
    body {{
        background: white !important;
    }}

    .resume-page {{
        width: 794px;
        min-height: 1123px;

        margin: 0;

        box-shadow: none;
    }}

    .resume-body {{
        break-inside: avoid;
    }}

    .main-section,
    .sidebar-section {{
        break-inside: avoid;
    }}
}}

</style>

</head>

<body>

<div class="resume-page">

    <!-- ================================================= -->
    <!-- HEADER -->
    <!-- ================================================= -->

    <header class="resume-header">

        <img
            class="avatar"
            src="{avatar_src}"
            alt="Profile photo"
        >

        <div>

            <h1
                class="m-0 text-[30px] font-extrabold
                       tracking-tight text-white"
            >
                {safe(name_display)}
            </h1>

            <div
                class="mt-2 text-[11px] font-bold
                       uppercase tracking-[0.18em]
                       text-sky-400"
            >
                {safe(title_display)}
            </div>

        </div>

    </header>


    <!-- ================================================= -->
    <!-- BODY -->
    <!-- ================================================= -->

    <div class="resume-body">

        <!-- SIDEBAR -->

        <aside class="sidebar">

            <section class="sidebar-section">

                <div class="section-title">
                    Contact
                </div>

                <div
                    class="mt-1 text-[12.5px]
                           leading-[1.55] text-slate-600"
                >
                    {contact_html}
                </div>

            </section>


            <section class="sidebar-section">

                <div class="section-title">
                    Languages
                </div>

                <div class="mt-1">
                    {language_html}
                </div>

            </section>


            <section class="sidebar-section">

                <div class="section-title">
                    Skills
                </div>

                <div class="mt-1">
                    {skill_html}
                </div>

            </section>

        </aside>


        <!-- MAIN -->

        <main class="main">

            <section class="main-section">

                <div class="main-title">
                    Summary
                </div>

                {summary_html}

            </section>


            <section class="main-section">

                <div class="main-title">
                    Education
                </div>

                {edu_html}

            </section>


            <section class="main-section">

                <div class="main-title">
                    Experience
                </div>

                {exp_html}

            </section>


            <section class="main-section">

                <div class="main-title">
                    Certifications
                </div>

                {cert_html}

            </section>

        </main>

    </div>

</div>

</body>

</html>
"""


# ============================================================
# STREAMLIT RESUME BUILDER
# ============================================================

def render_resume_builder(user_profile=None):
    if user_profile is None:
        user_profile = st.session_state.get("user_profile", {})
        st.set_page_config(
            page_title="Resume Builder",
            page_icon="📄",
            layout="wide",
    )

    # --------------------------------------------------------
    # STREAMLIT STYLING
    # --------------------------------------------------------

    st.markdown(
        """
        <style>

        .block-container {
            padding-top: 2rem;
            padding-bottom: 3rem;
            max-width: 1500px;
        }

        div[data-testid="stFileUploader"] {
            margin-bottom: 0.75rem;
        }

        .resume-preview-wrapper {
            background: #e2e8f0;
            border-radius: 16px;
            padding: 18px;
        }

        </style>
        """,
        unsafe_allow_html=True,
    )

    # --------------------------------------------------------
    # VALUE HELPER
    # --------------------------------------------------------

    def get_val(
        key,
        default="",
    ):

        value = user_profile.get(
            key,
            default,
        )

        if isinstance(value, list):
            return ", ".join(
                str(item)
                for item in value
            )

        if value is None:
            return ""

        return str(value)

    # --------------------------------------------------------
    # TWO COLUMNS
    # --------------------------------------------------------

    left, right = st.columns(
        [0.9, 1.1],
        gap="large",
    )

    # ========================================================
    # LEFT — FORM
    # ========================================================

    with left:

        st.title("📝 Resume Builder")

        st.caption(
            "Create a clean, professional resume "
            "and preview it live."
        )

        st.markdown("### Profile")

        uploaded_image = st.file_uploader(
            "Profile Picture",
            type=[
                "png",
                "jpg",
                "jpeg",
            ],
            help="Upload a professional profile photo.",
        )

        name = st.text_input(
            "Full Name",
            value=get_val(
                "name",
                "Galena Micheal",
            ),
        )

        professional_title = st.text_input(
            "Professional Title",
            value=get_val(
                "current_role",
                "BACHELOR OF ARTS IN EDUCATION",
            ),
        )

        st.markdown("### Contact Information")

        address = st.text_input(
            "Address",
            value=get_val(
                "address",
                "464 Canyon Trail, Charlotte, NC 48210, USA",
            ),
        )

        phone = st.text_input(
            "Phone",
            value=get_val(
                "phone",
                "(012) 444 6789",
            ),
        )

        email = st.text_input(
            "Email",
            value=get_val(
                "email",
                "gmicheal@email.com",
            ),
        )

        website = st.text_input(
            "Website",
            value=get_val(
                "website",
                "",
            ),
        )

        st.markdown("### Skills & Languages")

        languages_raw = st.text_input(
            "Languages",
            value=get_val(
                "languages",
                "English, Japanese, Spanish",
            ),
            help="Separate languages with commas.",
        )

        skills_raw = st.text_input(
            "Skills",
            value=get_val(
                "skills",
                "Organization, Leadership, Performance improvement plan development, Creative learning techniques, Behavior management",
            ),
            help="Separate skills with commas.",
        )

        languages = clean_list(
            languages_raw
        )

        skills = clean_list(
            skills_raw
        )

        st.markdown("### Professional Summary")

        summary_default = (
            "Motivated entry-level high school "
            "English Teacher with experience "
            "teaching multiple subject disciplines "
            "at all grade levels. Skilled in "
            "curriculum development, student "
            "performance improvement and "
            "classroom management."
        )

        summary = st.text_area(
            "Summary",
            value=get_val(
                "summary",
                summary_default,
            ),
            height=130,
        )

        st.markdown("### Education")

        edu_institution = st.text_input(
            "Institution",
            value=get_val(
                "edu_institution",
                "Johnson University",
            ),
        )

        edu_location = st.text_input(
            "Location",
            value=get_val(
                "edu_location",
                "West Charlotte, NC",
            ),
        )

        edu_degree = st.text_input(
            "Degree",
            value=get_val(
                "edu_degree",
                "Bachelor of Arts in Education",
            ),
        )

        edu_dates = st.text_input(
            "Dates",
            value=get_val(
                "edu_dates",
                "",
            ),
        )

        st.markdown("### Work Experience")

        exp_company = st.text_input(
            "Company / Organization",
            value=get_val(
                "exp_company",
                "River Tech High School",
            ),
        )

        exp_role = st.text_input(
            "Role",
            value=get_val(
                "exp_role",
                "Special Education Teacher",
            ),
        )

        exp_dates = st.text_input(
            "Dates",
            value=get_val(
                "exp_dates",
                "May 2023 – August 2024",
            ),
        )

        default_bullets = user_profile.get(
            "exp_bullets",
            [
                "Prepare 50+ students for the AP English Literature exam",
                "Discuss literary works, trends and techniques with students",
                "Administer written assignments and provide constructive feedback",
                "Create lesson plans and instructional resources for each class",
            ],
        )

        if isinstance(
            default_bullets,
            list,
        ):
            default_bullets = "\n".join(
                default_bullets
            )

        exp_bullets_raw = st.text_area(
            "Experience Bullet Points",
            value=str(
                default_bullets
            ),
            height=150,
            help="One achievement or responsibility per line.",
        )

        experience_bullets = [
            item.strip()
            for item in exp_bullets_raw.split(
                "\n"
            )
            if item.strip()
        ]

        st.markdown("### Certifications")

        default_certifications = user_profile.get(
            "certifications",
            [
                "Illinois Professional Educator License (PEL), 2025",
                "National Board Certified Teacher (NBCT), 2025",
            ],
        )

        if isinstance(
            default_certifications,
            list,
        ):
            default_certifications = "\n".join(
                default_certifications
            )

        cert_raw = st.text_area(
            "Certifications",
            value=str(
                default_certifications
            ),
            height=100,
            help="One certification per line.",
        )

        certifications = [
            item.strip()
            for item in cert_raw.split(
                "\n"
            )
            if item.strip()
        ]

    # ========================================================
    # IMAGE
    # ========================================================

    raw_img_bytes = None

    avatar_src = DEFAULT_AVATAR

    if uploaded_image is not None:

        raw_img_bytes = (
            uploaded_image.getvalue()
        )

        base64_img = base64.b64encode(
            raw_img_bytes
        ).decode("utf-8")

        mime_type = (
            uploaded_image.type
            or "image/png"
        )

        avatar_src = (
            f"data:{mime_type};base64,"
            f"{base64_img}"
        )

    # ========================================================
    # LIVE PROFILE
    # ========================================================

    live_profile = {
        "name": name,
        "current_role": professional_title,
        "address": address,
        "phone": phone,
        "email": email,
        "website": website,
        "summary": summary,
        "edu_institution": edu_institution,
        "edu_location": edu_location,
        "edu_degree": edu_degree,
        "edu_dates": edu_dates,
        "exp_company": exp_company,
        "exp_role": exp_role,
        "exp_dates": exp_dates,
        "exp_bullets": experience_bullets,
        "certifications": certifications,
        "skills": skills,
        "languages": languages,
    }

    # ========================================================
    # RIGHT — PREVIEW
    # ========================================================

    with right:

        st.title("👁️ Live Preview")

        st.caption(
            "The preview uses the same HTML layout "
            "that is used for the PDF."
        )

        full_html = build_resume_html(
            name=name,
            professional_title=professional_title,
            address=address,
            phone=phone,
            email=email,
            website=website,
            languages=languages,
            skills=skills,
            summary=summary,
            edu_institution=edu_institution,
            edu_location=edu_location,
            edu_degree=edu_degree,
            edu_dates=edu_dates,
            exp_company=exp_company,
            exp_role=exp_role,
            exp_dates=exp_dates,
            experience_bullets=experience_bullets,
            certifications=certifications,
            avatar_src=avatar_src,
        )

        # Important:
        # The previous version used 820px which cropped the
        # A4 document visually. This gives enough room to see
        # the entire first A4 page at once.
        st.components.v1.html(
            full_html,
            height=1180,
            scrolling=False,
        )

        st.markdown("---")

        st.subheader("📥 Download Resume")

        slug_raw = (
            name.strip().lower()
            if name.strip()
            else "resume"
        )

        file_name_slug = "".join(
            char
            if char.isalnum()
            else "_"
            for char in slug_raw
        )

        pdf_col, word_col = st.columns(
            2,
            gap="medium",
        )

        # ----------------------------------------------------
        # PDF
        # ----------------------------------------------------

        with pdf_col:

            if st.button(
                "📄 Generate PDF",
                use_container_width=True,
                type="primary",
            ):

                with st.spinner(
                    "Creating your PDF..."
                ):

                    pdf_bytes = (
                        generate_pdf_from_html(
                            full_html
                        )
                    )

                    if pdf_bytes is None:

                        pdf_bytes = (
                            generate_pdf_bytes(
                                live_profile,
                                raw_img_bytes,
                            )
                        )

                st.download_button(
                    "⬇️ Download PDF",
                    data=pdf_bytes,
                    file_name=(
                        f"{file_name_slug}"
                        "_resume.pdf"
                    ),
                    mime="application/pdf",
                    use_container_width=True,
                )

        # ----------------------------------------------------
        # WORD
        # ----------------------------------------------------

        with word_col:

            if st.button(
                "📝 Generate Word",
                use_container_width=True,
            ):

                with st.spinner(
                    "Creating your Word document..."
                ):

                    docx_bytes = (
                        generate_docx_bytes(
                            live_profile,
                            raw_img_bytes,
                        )
                    )

                st.download_button(
                    "⬇️ Download Word",
                    data=docx_bytes,
                    file_name=(
                        f"{file_name_slug}"
                        "_resume.docx"
                    ),
                    mime=(
                        "application/vnd.openxmlformats-"
                        "officedocument.wordprocessingml.document"
                    ),
                    use_container_width=True,
                )


# ============================================================
# RUN APP
# ============================================================

if __name__ == "__main__":
    render_resume_builder()